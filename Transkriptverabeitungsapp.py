import streamlit as st
from openai import OpenAI
import docx
import os
import io
import base64

# Bestimme den Pfad zum Logo
LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "logo.png")

def word_app():
    # Am Anfang der Funktion initialisieren wir den Session State
    if 'generated_responses' not in st.session_state:
        st.session_state.generated_responses = {}

    st.title("Chatte mit deinen Word-Dateien")

    # Eingabefeld für den OpenAI API Key (als Passwortfeld)
    openai_api_key = st.text_input("OpenAI API Key", type="password")

    # Vordefinierte Prompts
    predefined_prompts = {
        "Englische Übersetzung": "Du bist ein professioneller Übersetzer. Übersetze den folgenden Text ins Englische. Behalte dabei den Stil und Ton des Originaltextes bei. Die Übersetzung soll sich natürlich und professionell anhören.",
        "Zusammenfassung": "Erstelle eine prägnante Zusammenfassung des folgenden Textes. Behalte die wichtigsten Punkte bei und strukturiere die Zusammenfassung klar und übersichtlich. Die Zusammenfassung soll etwa 25% der Länge des Originaltextes betragen.",
        "Korrekturlesen": "Du bist ein professioneller Lektor. Überprüfe den Text auf Rechtschreibung, Grammatik und Stil. Verbessere den Text, wo nötig, und achte dabei auf eine klare und professionelle Ausdrucksweise.",
        "Benutzerdefiniert": "Gib hier dein Systemprompt ein..."
    }

    # Dropdown für vordefinierte Prompts
    selected_prompt = st.selectbox(
        "Wähle einen vordefinierten Prompt oder 'Benutzerdefiniert' für einen eigenen Prompt:",
        options=list(predefined_prompts.keys()),
        index=0  # Setzt "Englische Übersetzung" als Standard
    )

    # Eingabefeld für das Systemprompt
    system_prompt = st.text_area(
        "Systemprompt", 
        value=predefined_prompts[selected_prompt],
        height=150
    )

    # Mehrere Word-Dateien (docx) hochladen
    uploaded_files = st.file_uploader("Lade Word-Dateien hoch", type=["docx"], accept_multiple_files=True)

    # Prüfen, ob API Key, Systemprompt und Dateien vorhanden sind
    if openai_api_key and system_prompt and uploaded_files:
        # Initialisiere den OpenAI-Client mit dem API-Key
        client = OpenAI(api_key=openai_api_key)
        st.write("### Verarbeitung der Dateien:")
        
        # Durchlaufe alle hochgeladenen Dateien
        for uploaded_file in uploaded_files:
            st.write(f"**Datei:** {uploaded_file.name}")
            
            # Lese den gesamten Text aus der Word-Datei
            try:
                doc = docx.Document(uploaded_file)
                full_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                st.error(f"Fehler beim Lesen der Datei {uploaded_file.name}: {e}")
                continue
            
            # Zeige den extrahierten Text in einem Textfeld an
            st.text_area(f"Extrahierter Text aus {uploaded_file.name}:", full_text, height=200)
            
            # Erstelle die Nachrichten für den GPT-4 Chat
            try:
                # Teile den Text in größere Chunks (ca. 32k Tokens ≈ 40000 Zeichen)
                chunk_size = 40000
                text_chunks = [full_text[i:i+chunk_size] 
                             for i in range(0, len(full_text), chunk_size)]
                
                complete_output = []
                
                # Erstelle einen Chat-Container für die Ausgabe
                with st.chat_message("assistant", avatar=LOGO_PATH):
                    # Erstelle einen leeren Platz für den Stream
                    message_placeholder = st.empty()
                    # Sammle den vollständigen Text
                    full_response = ""
                    
                    # Wenn es sich um eine Zusammenfassung handelt, den gesamten Text auf einmal verarbeiten
                    if "Zusammenfassung" in system_prompt:
                        messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": full_text}
                        ]
                        
                        # Stream die Antwort für den gesamten Text
                        response_stream = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=messages,
                            temperature=0.2,
                            max_tokens=16000,
                            stream=True
                        )
                        
                        # Verarbeite den Stream
                        for response in response_stream:
                            if response.choices[0].delta.content is not None:
                                full_response += response.choices[0].delta.content
                                message_placeholder.markdown(full_response + "▌")
                            
                    else:
                        # Originale Chunk-Verarbeitung für andere Prompts
                        for chunk_index, chunk in enumerate(text_chunks):
                            messages = [
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": chunk}
                            ]
                            
                            # Wenn es nicht der erste Chunk ist, füge Kontext hinzu
                            if complete_output:
                                messages.append({
                                    "role": "assistant",
                                    "content": "Bisheriger Kontext: " + " ".join(complete_output)
                                })
                            
                            # Stream die Antwort
                            response_stream = client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=messages,
                                temperature=0.2,
                                max_tokens=16000,
                                stream=True  # Aktiviere Streaming
                            )
                            
                            # Chunk-Header anzeigen
                            if len(text_chunks) > 1:
                                full_response += f"\n\nTeil {chunk_index + 1} von {len(text_chunks)}:\n"
                            
                            # Verarbeite den Stream
                            for response in response_stream:
                                if response.choices[0].delta.content is not None:
                                    full_response += response.choices[0].delta.content
                                    # Aktualisiere die Anzeige in Echtzeit
                                    message_placeholder.markdown(full_response + "▌")
                            
                            complete_output.append(full_response)
                            
                        # Finale Anzeige ohne Cursor
                        message_placeholder.markdown(full_response)
                
                # Speichere die generierte Antwort im Session State
                st.session_state.generated_responses[uploaded_file.name] = full_response

                # Erstelle Word-Datei
                output_doc = docx.Document()
                output_doc.add_paragraph(st.session_state.generated_responses[uploaded_file.name])
                doc_io = io.BytesIO()
                output_doc.save(doc_io)
                doc_io.seek(0)
                
                # Konvertiere zu base64
                docx_b64 = base64.b64encode(doc_io.getvalue()).decode()
                
                # Erstelle einen HTML Download-Link
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_b64}" download="{uploaded_file.name}_output.docx">Download als DOCX</a>'
                st.markdown(href, unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"Fehler bei der Verarbeitung durch GPT-4 für {uploaded_file.name}: {e}")
    else:
        st.info("Bitte gib deinen OpenAI API Key, das Systemprompt ein und lade mindestens eine Word-Datei hoch.")

if __name__ == "__main__":
    word_app()
