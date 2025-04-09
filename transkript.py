import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import streamlit as st
import replicate
import requests
import tempfile
import smtplib
from email.message import EmailMessage
import math
import io
import base64

def send_email_notification(receiver_emails, file_name, download_link):
    sender_email = "jonathan.heeckt@bonsai-research.com" 
    password = "pjtpqvtmdkrhfgvk"  

    # Send the email to each receiver
    for receiver_email in receiver_emails:
        msg = EmailMessage()
        msg['Subject'] = 'Transkription abgeschlossen'
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg.set_content(f"Moin,\n\nDie Transkription für '{file_name}' ist abgeschlossen und kann hier heruntergeladen werden: {download_link}\n\nLG Jonathan")

        with smtplib.SMTP('smtp.office365.com', 587) as smtp:
            smtp.starttls()
            smtp.login(sender_email, password)
            smtp.send_message(msg)


# Constants for styles and files
BASE_STYLE = 'Normal'
FONT_NAME = 'UCity'
FONT_SIZE = 12
SPEAKER_ZERO_COLOR = (245, 0, 126)     # Bonsai Pink
SPEAKER_ONE_COLOR = (0, 0, 0)          # Black
SPEAKER_TWO_COLOR = (0, 0, 255)        # Blue
SPEAKER_THREE_COLOR = (0, 255, 0)      # Green
SPEAKER_FOUR_COLOR = (255, 0, 0)       # Red
SPEAKER_FIVE_COLOR = (255, 255, 0)     # Yellow
SPEAKER_SIX_COLOR = (255, 0, 255)      # Magenta
SPEAKER_SEVEN_COLOR = (0, 255, 255)    # Cyan
SPEAKER_EIGHT_COLOR = (128, 0, 0)      # Maroon
SPEAKER_NINE_COLOR = (0, 128, 0)       # Dark Green
SPEAKER_TEN_COLOR = (0, 0, 128)        # Navy
TRANSCRIPT_FILENAME = "Transkript.docx"

class StyleError(Exception):
    pass

class SegmentKeyError(Exception):
    pass

class TimestampError(Exception):
    pass

class AudioProcessError(Exception):
    pass

def get_styles(doc, name, font_size=FONT_SIZE, color_rgb=(0, 0, 0), base_style='Normal', font_name='UCity'):
    """Creates a new Word style with provided parameters."""
    try:
        style = doc.styles.add_style(name, doc.styles[base_style].type)
        
        # Explicitly set the font to UCity for this style
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(*color_rgb)
        style.paragraph_format.left_indent = Inches(0.1)
        style.paragraph_format.right_indent = Inches(0.1)
        style.paragraph_format.space_after = Inches(0.1)
        style.paragraph_format.space_before = Inches(0.0)
        
        return style
    except Exception as e:
        raise StyleError(f'Error while generating style: {e}')


def format_json_to_chat(segments, output_file_path, speaker_names=None):   
    """Transforms a JSON object into a formatted conversation saved as a Word document."""
    doc = Document()
    font = doc.styles['Normal'].font
    font.name = 'UCity'

    # Get styles for each speaker
    speaker_styles = {
        "SPEAKER_00": get_styles(doc, 'Speaker0', color_rgb=SPEAKER_ZERO_COLOR),
        "SPEAKER_01": get_styles(doc, 'Speaker1', color_rgb=SPEAKER_ONE_COLOR),
        "SPEAKER_02": get_styles(doc, 'Speaker2', color_rgb=SPEAKER_TWO_COLOR),
        "SPEAKER_03": get_styles(doc, 'Speaker3', color_rgb=SPEAKER_THREE_COLOR),
        "SPEAKER_04": get_styles(doc, 'Speaker4', color_rgb=SPEAKER_FOUR_COLOR),
        "SPEAKER_05": get_styles(doc, 'Speaker5', color_rgb=SPEAKER_FIVE_COLOR),
        "SPEAKER_06": get_styles(doc, 'Speaker6', color_rgb=SPEAKER_SIX_COLOR),
        "SPEAKER_07": get_styles(doc, 'Speaker7', color_rgb=SPEAKER_SEVEN_COLOR),
        "SPEAKER_08": get_styles(doc, 'Speaker8', color_rgb=SPEAKER_EIGHT_COLOR),
        "SPEAKER_09": get_styles(doc, 'Speaker9', color_rgb=SPEAKER_NINE_COLOR),
        "SPEAKER_10": get_styles(doc, 'Speaker10', color_rgb=SPEAKER_TEN_COLOR)
    }

    last_speaker = None

    for segment in segments:
        try: 
            original_speaker_label = segment['speaker']  # Save the original label
            if speaker_names and original_speaker_label in speaker_names:
                display_speaker_name = speaker_names[original_speaker_label]  # This is the name to be displayed in the doc
            else:
                display_speaker_name = original_speaker_label

            text = segment['text']
            formatted_timestamp = format_segment_time(segment)
            
            # If speaker changes, add an additional line break
            if last_speaker and last_speaker != display_speaker_name:
                doc.add_paragraph()

            # Add speaker's name and timestamp in bold
            p = doc.add_paragraph()
            p.add_run(f"{display_speaker_name} {formatted_timestamp}: ").bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add speaker's text on the next line
            p = doc.add_paragraph(text, style=speaker_styles.get(original_speaker_label, 'Normal'))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            last_speaker = display_speaker_name
        except KeyError as e:
            raise SegmentKeyError(f'Missing key in segment: {e}')

    doc.save(output_file_path)

def format_segment_time(segment):
    """Transforms a start and end time into timestamp format [MM:SS-MM:SS]."""
    try:
        start_time = float(segment['start'])
        end_time = float(segment['end'])
        start_minutes, start_seconds = divmod(int(start_time), 60)
        end_minutes, end_seconds = divmod(int(end_time), 60)
        return f"[{start_minutes:02}:{start_seconds:02}-{end_minutes:02}:{end_seconds:02}]"
    except ValueError as e:
        raise TimestampError(f'Error while parsing time values: {e}')
    
def upload_to_hosting_service(uploaded_file):
    """
    Uploads the file to tmpfiles.org and returns the direct URL to the file.
    """
    # The API endpoint for uploading files
    upload_url = "https://tmpfiles.org/api/v1/upload"
    
    # Post the file to the endpoint
    response = requests.post(upload_url, files={'file': uploaded_file})
    
    # Check if the upload was successful
    response.raise_for_status()
    
    # Parse the JSON response to get the direct file URL
    response_data = response.json()
    file_url = response_data.get('data', {}).get('url')
    
    if not file_url:
        raise ValueError(f"Failed to retrieve the file URL from the tmpfiles response. Full response: {response_data}")
    
    # Construct the direct download link
    direct_download_link = file_url.replace("https://tmpfiles.org/", "https://tmpfiles.org/dl/")
    
    return direct_download_link


def process_audio(uploaded_file=None, direct_url=None, num_speakers=2, language="", prompt="", base_name=None, speaker_names=None, translate=False):
    try:
        # Erstelle einen expliziten Client mit dem Token
        if 'replicate_token' in st.session_state:
            client = replicate.Client(api_token=st.session_state.replicate_token)
            
        if direct_url:  # If a direct URL is provided
            file_url = direct_url
        elif uploaded_file:  # If an audio file is uploaded
            # Save the uploaded file to a temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio:
                temp_audio.write(uploaded_file.read())
                file_url = upload_to_hosting_service(open(temp_audio.name, 'rb'))
                
                # Print the file URL to verify
                print(f"Uploaded file URL: {file_url}")
                
        else:
            raise ValueError("Either a direct URL or an uploaded file is required.")

        # Verwende den Client anstelle von replicate.run
        output = client.run(
            "thomasmol/whisper-diarization:cbd15da9f839c5f932742f86ce7def3a03c22e2b4171d42823e83e314547003f",
            input={
                "file_url": file_url,
                "group_segments": True,
                "num_speakers": num_speakers,
                "language": language,
                "prompt": prompt,
                "offset_seconds": 0,
                "translate": translate,
            }
        )
        
        if base_name:
            output_filename = f"{base_name}.docx"
        else:
            output_filename = TRANSCRIPT_FILENAME

        if 'segments' in output:
            output_path = os.path.join(os.getcwd(), output_filename)
            format_json_to_chat(output['segments'], output_path, speaker_names)

            # Upload the transcript file and get the download link
            with open(output_path, 'rb') as transcript_file:
                transcript_download_link = upload_to_hosting_service(transcript_file)

            return output_path, transcript_download_link  # Return both the local path and download link
        else:
            raise AudioProcessError('The output is not as expected')
    except Exception as e:
        raise AudioProcessError(f'Error in processing audio: {e}')



def main():
    st.title("BonsAI Transkriptionsapp")
    if not enter_replicate_api_token():
        st.warning("⚠️ Bitte gib einen gültigen API Token ein, um fortzufahren.")
        st.stop()

    # Eingabefeld für die Anzahl der Transkripte
    num_transcripts = st.number_input("Anzahl der Transkripte:", min_value=1, value=1)

    # Container für jedes Transkript erstellen
    for transcript_idx in range(num_transcripts):
        with st.expander(f"Transkript {transcript_idx + 1}", expanded=True):
            # Zwei Spalten erstellen für Dropdown und manuellen Input
            col1, col2 = st.columns(2)
            
            with col1:
                language_dropdown = st.selectbox(
                    f"Welche Sprache wird gesprochen? (Transkript {transcript_idx + 1})",
                    ('de', 'en', 'it', 'fr', 'es', 'pl', 'sv', 'dk', 'bu', 'nl', 'hu'),
                    index=None,
                    placeholder="Bitte wähle die Sprache",
                    key=f"lang_dropdown_{transcript_idx}"
                )
            
            with col2:
                language_input = st.text_input(
                    f"Oder manuelle Eingabe des Sprachcodes (Transkript {transcript_idx + 1})",
                    "",
                    key=f"lang_input_{transcript_idx}"
                )
            
            # Bestimme die zu verwendende Sprache
            language = language_input if language_input else language_dropdown
            
            # Anzahl der Sprecher
            num_speakers = st.number_input(
                f"Anzahl der Sprecher (Transkript {transcript_idx + 1}):", 
                min_value=1, 
                max_value=10, 
                value=2,
                key=f"num_speakers_{transcript_idx}"
            )
            
            # Prompt für zusätzlichen Kontext
            prompt = st.text_area(
                f"Zusätzlicher Kontext (optional) (Transkript {transcript_idx + 1}):", 
                "", 
                key=f"prompt_{transcript_idx}"
            )
            
            # Sprechernamen zuweisen
            with st.expander(f"Sprechernamen zuweisen (Transkript {transcript_idx + 1})", expanded=False):
                speaker_names = {}
                for speaker_idx in range(num_speakers):
                    speaker_label = f"SPEAKER_{speaker_idx:02d}"
                    speaker_name = st.text_input(
                        f"Name für {speaker_label}", 
                        "", 
                        key=f"speaker_name_{transcript_idx}_{speaker_idx}"
                    )
                    if speaker_name:
                        speaker_names[speaker_label] = speaker_name
            
            # Übersetzungsoption
            translate_to_english = st.checkbox(
                f"Ins Englische übersetzen (Transkript {transcript_idx + 1})", 
                False,
                key=f"translate_{transcript_idx}"
            )
            
            # Datei hochladen
            uploaded_file = st.file_uploader(
                f"Audio- oder Videodatei hochladen (Transkript {transcript_idx + 1})", 
                type=["mp3", "mp4", "wav", "m4a", "webm"],
                key=f"file_uploader_{transcript_idx}"
            )
            
            # Alternativ URL eingeben
            direct_url = st.text_input(
                f"Oder direkte URL zur Datei (Transkript {transcript_idx + 1})", 
                "",
                key=f"direct_url_{transcript_idx}"
            )
            
            # Basisname für die Ausgabedatei
            base_name = st.text_input(
                f"Ausgabedateiname (ohne Erweiterung) (Transkript {transcript_idx + 1})", 
                f"Transkript_{transcript_idx + 1}",
                key=f"base_name_{transcript_idx}"
            )
            
            # E-Mail-Adresse für Benachrichtigung
            email = st.text_input(
                f"E-Mail-Adresse für Benachrichtigung (Transkript {transcript_idx + 1})", 
                "",
                key=f"email_{transcript_idx}"
            )
            emails = [email] if email else []
    
    # Button zum Starten des Prozesses
    if st.button("Transkription starten"):
        successful = True
        # Verarbeite jedes Transkript
        for transcript_idx in range(num_transcripts):
            # Extrahiere die Parameter für das aktuelle Transkript
            language_dropdown = st.session_state.get(f"lang_dropdown_{transcript_idx}")
            language_input = st.session_state.get(f"lang_input_{transcript_idx}")
            language = language_input if language_input else language_dropdown
            
            num_speakers = st.session_state.get(f"num_speakers_{transcript_idx}")
            prompt = st.session_state.get(f"prompt_{transcript_idx}")
            translate = st.session_state.get(f"translate_{transcript_idx}")
            uploaded_file = st.session_state.get(f"file_uploader_{transcript_idx}")
            direct_url = st.session_state.get(f"direct_url_{transcript_idx}")
            base_name = st.session_state.get(f"base_name_{transcript_idx}")
            email = st.session_state.get(f"email_{transcript_idx}")
            emails = [email] if email else []
            
            # Extrahiere Sprechernamen
            speaker_names = {}
            for speaker_idx in range(num_speakers):
                speaker_label = f"SPEAKER_{speaker_idx:02d}"
                speaker_name = st.session_state.get(f"speaker_name_{transcript_idx}_{speaker_idx}")
                if speaker_name:
                    speaker_names[speaker_label] = speaker_name
            
            # Verarbeite das Transkript
            if (uploaded_file or direct_url) and language:
                try:
                    # Verarbeite das Transkript
                    result = handle_audio_process(
                        num_speakers=num_speakers,
                        language=language,
                        prompt=prompt,
                        user_email=emails,
                        uploaded_file=uploaded_file,
                        direct_url=direct_url,
                        base_name=base_name,
                        speaker_names=speaker_names,
                        translate=translate
                    )
                    
                    if result:
                        st.success(f"✅ Transkription {transcript_idx + 1} erfolgreich! {result}")
                    else:
                        st.error(f"❌ Fehler bei der Verarbeitung von Transkript {transcript_idx + 1}!")
                        successful = False
                except Exception as e:
                    st.error(f"❌ Fehler bei der Verarbeitung von Transkript {transcript_idx + 1}: {str(e)}")
                    successful = False
            else:
                st.error(f"❌ Bitte lade eine Datei hoch oder gib eine URL ein und wähle eine Sprache für Transkript {transcript_idx + 1}!")
                successful = False
        
        # Gesamtergebnis anzeigen
        if successful:
            st.balloons()
            st.success("✅ Alle Transkripte wurden erfolgreich verarbeitet!")
            st.info("ℹ️ Wenn du eine E-Mail-Adresse angegeben hast, erhältst du eine Benachrichtigung, sobald die Transkription vollständig ist.")
        else:
            st.warning("⚠️ Es sind Fehler bei der Verarbeitung aufgetreten. Bitte überprüfe die Fehlermeldungen oben.")

def enter_replicate_api_token():
    """
    Erlaubt dem Benutzer, seinen Replicate API-Token einzugeben und speichert ihn im Session State.
    Gibt True zurück, wenn ein Token im Session State existiert, sonst False.
    """
    if 'replicate_token' not in st.session_state:
        st.session_state.replicate_token = ""
        
    # Token-Eingabe
    token = st.text_input(
        "Replicate API Token",
        value=st.session_state.replicate_token,
        type="password",
        help="Du benötigst einen Replicate API Token. Erhalte ihn unter https://replicate.com/account/api-tokens"
    )
    
    # Token speichern
    if token:
        st.session_state.replicate_token = token
        return True
    return False

def handle_audio_process(num_speakers, language, prompt, user_email, uploaded_file=None, direct_url=None, base_name=None, speaker_names=None, translate=False):
    """
    Verarbeitet eine Audiodatei und gibt eine Erfolgsmeldung oder eine Fehlermeldung zurück.
    """
    try:
        with st.spinner(f"Transkription läuft... (Sprache: {language}, Sprecher: {num_speakers})"):
            # Dateigröße prüfen und ggf. aufteilen
            if uploaded_file:
                file_size = uploaded_file.size
                # Wenn die Datei größer als 95 MB ist, aufteilen
                if file_size > 95 * 1024 * 1024:  # 95 MB in Bytes
                    st.warning(f"Die Datei ist {file_size / (1024 * 1024):.2f} MB groß und wird aufgeteilt.")
                    file_chunks = split_file(uploaded_file)
                    st.info(f"Datei in {len(file_chunks)} Teile aufgeteilt.")
                    
                    # Jedes Stück verarbeiten
                    for i, chunk in enumerate(file_chunks):
                        output_path, download_link = process_audio(
                            uploaded_file=chunk,
                            num_speakers=num_speakers,
                            language=language,
                            prompt=prompt,
                            base_name=f"{base_name}_Teil_{i+1}" if base_name else f"Transkript_Teil_{i+1}",
                            speaker_names=speaker_names,
                            translate=translate
                        )
                        
                        # E-Mail-Benachrichtigung senden, falls gewünscht
                        if user_email:
                            send_email_notification(user_email, f"{base_name}_Teil_{i+1}" if base_name else f"Transkript_Teil_{i+1}", download_link)
                        
                        st.success(f"Teil {i+1}/{len(file_chunks)} verarbeitet. [Download]({download_link})")
                        
                    return "Alle Teile wurden verarbeitet. Siehe Links oben zum Herunterladen."
            
            # Standard-Verarbeitung für Dateien unter 95 MB oder direkte URLs
            output_path, download_link = process_audio(
                uploaded_file=uploaded_file,
                direct_url=direct_url,
                num_speakers=num_speakers,
                language=language,
                prompt=prompt,
                base_name=base_name,
                speaker_names=speaker_names,
                translate=translate
            )
            
            # E-Mail-Benachrichtigung senden, falls gewünscht
            if user_email:
                send_email_notification(user_email, base_name or "Transkript", download_link)
                
            # Download-Button anzeigen
            return f"[Transkript herunterladen]({download_link})"
    except Exception as e:
        st.error(f"Fehler während der Verarbeitung: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

def split_file(file, max_size=95 * 1024 * 1024):
    """
    Teilt eine Datei in Stücke von max_size auf und gibt die Stücke als BytesIO-Objekte zurück.
    """
    # Datei komplett in den Speicher lesen
    file_content = file.read()
    file_size = len(file_content)
    
    # Anzahl der benötigten Stücke berechnen
    num_chunks = math.ceil(file_size / max_size)
    chunks = []
    
    for i in range(num_chunks):
        start = i * max_size
        end = min((i + 1) * max_size, file_size)
        chunk_content = file_content[start:end]
        
        # BytesIO-Objekt erstellen
        chunk = io.BytesIO(chunk_content)
        chunk.name = f"{file.name}.part{i+1}"  # Name für das Stück setzen
        chunks.append(chunk)
    
    return chunks

if __name__ == "__main__":
    main()
